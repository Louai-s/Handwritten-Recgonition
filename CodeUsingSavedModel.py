#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jul 26 01:53:00 2021

@author: louaisalem
"""

from tensorflow import keras
import numpy as np
import pandas as pd
from keras.preprocessing.image import ImageDataGenerator
import os
import random 
import cv2
import imutils
import random
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import LabelBinarizer
from keras.utils import np_utils
from keras.models import Sequential
from keras import optimizers
from keras import backend as K
from keras.layers import Dense, Activation, Flatten, Dense,MaxPooling2D, Dropout
from keras.layers import Conv2D, MaxPooling2D, BatchNormalization
from scipy.spatial import distance as dist

#Load Data to the Test array from The file    
Test_dir="Data/A-Z_Char/Test"
test_data=[]
img_size=32
non_chars = ["#","$","&","@"]
for i in os.listdir(Test_dir):
    if i in non_chars:
        continue
    count = 0
    if not i.startswith('.'):
        sub_directory = os.path.join(Test_dir,i)
        for j in os.listdir(sub_directory):
            if not j.startswith('.'):
                count+=1
                if count > 500:
                    break
                img = cv2.imread(os.path.join(sub_directory,j),0)
                img = cv2.resize(img,(img_size,img_size))
                test_data.append([img,i])
                
test_X=[]
test_Y=[]
for features,label in test_data:
    test_X.append(features)
    test_Y.append(label)

LB = LabelBinarizer()
test_Y=LB.fit_transform(test_Y)
test_X= np.array(test_X)/255.0
test_X=test_X.reshape(-1,32,32,1)
test_Y=np.array(test_Y)

#-------------------------------------------------Loading The Saved Model And evaluate-------------------------------------------------
SavedModel=keras.models.load_model("Code/TheModelNew.h5")

SavedModel.evaluate(test_X,test_Y)

Predictions=SavedModel.predict(test_X) # predict test images
Predictions=LB.inverse_transform(Predictions) #convert the prediction from numeric to the real labels (not numeric)

ActualLabels=LB.inverse_transform(test_Y) #convert the actual labels from numeric to the Categorical labels
from sklearn.metrics import confusion_matrix,f1_score

from sklearn.metrics import classification_report
indexes = np.unique(ActualLabels, return_index=True)[1] # Get The labels without duplicates
labels=[ActualLabels[index] for index in sorted(indexes)] # Keep their order like it was in the array
labels=np.sort(labels)
cm=confusion_matrix(ActualLabels,Predictions) # build the confusion matrix
print(classification_report(ActualLabels, Predictions))

#Plot The Confusion matrix
from mlxtend.plotting import plot_confusion_matrix
font={
      'family':'New Times Roman',
      'size':20
      }
plt.rc('font', **font)
plot_confusion_matrix(conf_mat=cm,figsize=(25,25),class_names=labels)

#-------------------------------------------------Post-Processing-------------------------------------------------
def sort_contours(cnts, method="left-to-right"):
    reverse = False
    i = 0
    if method == "right-to-left" or method == "bottom-to-top":
        reverse = True
    if method == "top-to-bottom" or method == "bottom-to-top":
        i = 1
    boundingBoxes = [cv2.boundingRect(c) for c in cnts]
    (cnts, boundingBoxes) = zip(*sorted(zip(cnts, boundingBoxes),
    key=lambda b:b[1][i], reverse=reverse))
    # return the list of sorted contours and bounding boxes
    return (cnts, boundingBoxes)

import math
def dist(a,b):
    return math.sqrt((a[0]-b[0])**2 + (a[1]-b[1])**2)

def get_letters(img):
    letters = []
    image = cv2.imread(img)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    ret,thresh1 = cv2.threshold(gray ,127,255,cv2.THRESH_BINARY_INV)
    dilated = cv2.dilate(thresh1, None, iterations=2)

    cnts = cv2.findContours(dilated.copy(), cv2.RETR_EXTERNAL,cv2.CHAIN_APPROX_SIMPLE)
    cnts = imutils.grab_contours(cnts)
    cnts = sort_contours(cnts, method="left-to-right")[0]
    c1=0
    Centers=[]
    # loop over the contours
    for c in cnts:
        if cv2.contourArea(c) > 10:
            (x, y, w, h) = cv2.boundingRect(c)
            cv2.rectangle(image, (x, y), (x + w, y + h), (0, 255, 0), 2)
        CenterX=x+0.5*w
        CenterY=y+0.5*h
        CenterXY=[CenterX,CenterY]
        Centers.append(CenterXY)
        if c1>0:
            if dist(Centers[c1-1], Centers[c1])>150:
                letters.append(" ")
        c1=c1+1
        roi = gray[y:y + h, x:x + w]
        thresh = cv2.threshold(roi, 127, 255,cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
        thresh = cv2.resize(thresh, (32, 32), interpolation = cv2.INTER_CUBIC)
        thresh = thresh.astype("float32") / 255.0
        thresh = np.expand_dims(thresh, axis=-1)
        thresh = thresh.reshape(1,32,32,1)
        ypred = SavedModel.predict(thresh)
        ypred = LB.inverse_transform(ypred)
        [x] = ypred
        letters.append(x)
    return letters, image

    plt.imshow(image)    
    

def get_word(letter):
    word = "".join(letter)
    return word

#-------------------------------------------------Testing The model/Making Some Predictions-------------------------------------------------
# The Predict method included in get_letters function
letters,image = get_letters("Data/ILoveAI.png")
word = get_word(letters) #To get the word without ","
print(word)
plt.imshow(image)

#Printing The word to The word document
from docx import Document
from docx.shared import Pt
import os
doc=Document()
parag=doc.add_paragraph(word)
font=parag.style.font
font.size=Pt(35)
font.bold=True
doc.save("test.docx")

#Making another one
letters,image = get_letters("Data/LouaiAwni.png")
word = get_word(letters)
print(word)
plt.imshow(image)

parag=doc.add_paragraph(word)
font=parag.style.font
font.size=Pt(35)
font.bold=True
doc.save("test.docx")
os.system("open test.docx")


